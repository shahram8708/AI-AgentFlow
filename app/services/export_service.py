"""Export generation service."""

from __future__ import annotations

import csv
import html
import io
import json
import re
from datetime import datetime
from io import BytesIO
from typing import Any

import pytz
from docx import Document
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Pt, RGBColor
from reportlab.lib import colors
from reportlab.lib.pagesizes import A4
from reportlab.lib.styles import ParagraphStyle, getSampleStyleSheet
from reportlab.lib.units import inch
from reportlab.platypus import ListFlowable, ListItem, Paragraph, Preformatted, SimpleDocTemplate, Spacer

from app.models import AutomationTask, TaskOutput
from app.services.file_service import FileService
from app.utils.sanitizer import sanitize_filename


class ExportServiceError(Exception):
    """Raised when export generation fails."""


class ExportService:
    """Exports task outputs to DOCX, PDF, CSV, and JSON."""

    def __init__(self) -> None:
        self.file_service = FileService()

    @staticmethod
    def _to_ist_string(value: datetime | None) -> str:
        if value is None:
            return ""
        ist = pytz.timezone("Asia/Kolkata")
        if value.tzinfo is None:
            value = pytz.utc.localize(value)
        return value.astimezone(ist).strftime("%d %b %Y, %I:%M %p IST")

    @staticmethod
    def _task_name_slug(task_name: str) -> str:
        normalized = re.sub(r"\s+", "_", task_name.strip().lower())
        normalized = sanitize_filename(normalized)
        if not normalized:
            normalized = "agentflow_output"
        return normalized[:50]

    @staticmethod
    def _parse_markdown_lines(content: str) -> list[str]:
        return content.splitlines() if content else []

    @staticmethod
    def _strip_line_markdown_prefix(line: str) -> str:
        stripped = line.strip()
        if stripped.startswith("### "):
            return stripped[4:].strip()
        if stripped.startswith("## "):
            return stripped[3:].strip()
        if stripped.startswith("# "):
            return stripped[2:].strip()
        if stripped.startswith("- ") or stripped.startswith("* "):
            return stripped[2:].strip()
        if re.match(r"^\d+\.\s+", stripped):
            return re.sub(r"^\d+\.\s+", "", stripped)
        return line

    @staticmethod
    def _iter_inline_markdown_tokens(text: str) -> list[dict[str, Any]]:
        token_pattern = re.compile(
            r"(\*\*[^*]+\*\*|__[^_]+__|\*[^*]+\*|_[^_]+_|`[^`]+`|\[[^\]]+\]\([^)]+\))"
        )
        tokens: list[dict[str, Any]] = []
        last_idx = 0

        for match in token_pattern.finditer(text):
            start, end = match.span()
            if start > last_idx:
                tokens.append({"text": text[last_idx:start], "emphasis": False, "code": False})

            raw = match.group(0)
            if raw.startswith("**") and raw.endswith("**"):
                tokens.append({"text": raw[2:-2], "emphasis": True, "code": False})
            elif raw.startswith("__") and raw.endswith("__"):
                tokens.append({"text": raw[2:-2], "emphasis": True, "code": False})
            elif raw.startswith("*") and raw.endswith("*"):
                tokens.append({"text": raw[1:-1], "emphasis": True, "code": False})
            elif raw.startswith("_") and raw.endswith("_"):
                tokens.append({"text": raw[1:-1], "emphasis": True, "code": False})
            elif raw.startswith("`") and raw.endswith("`"):
                tokens.append({"text": raw[1:-1], "emphasis": False, "code": True})
            else:
                link_match = re.match(r"^\[([^\]]+)\]\(([^)]+)\)$", raw)
                if link_match:
                    tokens.append(
                        {
                            "text": f"{link_match.group(1)} ({link_match.group(2)})",
                            "emphasis": False,
                            "code": False,
                        }
                    )
                else:
                    tokens.append({"text": raw, "emphasis": False, "code": False})

            last_idx = end

        if last_idx < len(text):
            tokens.append({"text": text[last_idx:], "emphasis": False, "code": False})

        return [token for token in tokens if token["text"]]

    @classmethod
    def _to_plain_markdown_text(cls, text: str) -> str:
        if not text:
            return ""

        output_lines: list[str] = []
        for raw_line in text.splitlines():
            line = cls._strip_line_markdown_prefix(raw_line)
            tokens = cls._iter_inline_markdown_tokens(line)
            output_lines.append("".join(token["text"] for token in tokens))
        return "\n".join(output_lines)

    def _write_docx_markdown_runs(self, paragraph: Any, text: str) -> None:
        dark_text = RGBColor(17, 24, 39)
        for token in self._iter_inline_markdown_tokens(text):
            run = paragraph.add_run(token["text"])
            if token["emphasis"]:
                run.bold = True
                run.font.color.rgb = dark_text
            if token["code"]:
                run.font.name = "Courier New"
                run.font.size = Pt(10)

    def _to_pdf_paragraph_markup(self, text: str) -> str:
        chunks: list[str] = []
        for token in self._iter_inline_markdown_tokens(text):
            escaped = html.escape(token["text"])
            if token["emphasis"]:
                chunks.append(f'<font color="#111827"><b>{escaped}</b></font>')
            elif token["code"]:
                chunks.append(f"<font name=\"Courier\">{escaped}</font>")
            else:
                chunks.append(escaped)
        return "".join(chunks)

    @staticmethod
    def _parse_markdown_heading(line: str) -> tuple[int, str] | None:
        match = re.match(r"^(#{1,6})\s+(.+)$", line.strip())
        if not match:
            return None
        return len(match.group(1)), match.group(2).strip()

    @staticmethod
    def _is_numbered_list(line: str) -> bool:
        return bool(re.match(r"^\d+\.\s+", line.strip()))

    @staticmethod
    def _add_docx_bottom_border(paragraph: Any) -> None:
        p_pr = paragraph._p.get_or_add_pPr()
        p_borders = OxmlElement("w:pBdr")
        p_bottom = OxmlElement("w:bottom")
        p_bottom.set(qn("w:val"), "single")
        p_bottom.set(qn("w:sz"), "6")
        p_bottom.set(qn("w:space"), "1")
        p_bottom.set(qn("w:color"), "D1D5DB")
        p_borders.append(p_bottom)
        p_pr.append(p_borders)

    @staticmethod
    def _add_docx_page_number(paragraph: Any) -> None:
        paragraph.add_run("Page ")
        fld = OxmlElement("w:fldSimple")
        fld.set(qn("w:instr"), "PAGE")
        paragraph._p.append(fld)

    def export_to_docx(self, task: AutomationTask, output: TaskOutput) -> BytesIO:
        """Export task output to a professionally formatted DOCX document."""

        document = Document()

        normal_style = document.styles["Normal"]
        normal_style.font.name = "Calibri"
        normal_style.font.size = Pt(11)

        heading1 = document.styles["Heading 1"]
        heading1.font.color.rgb = RGBColor(15, 23, 42)
        heading2 = document.styles["Heading 2"]
        heading2.font.color.rgb = RGBColor(30, 64, 175)
        heading3 = document.styles["Heading 3"]
        heading3.font.color.rgb = RGBColor(3, 105, 161)

        title = document.add_heading(task.task_name or "AgentFlow Output", level=0)
        title_run = title.runs[0]
        title_run.font.name = "Calibri"
        title_run.font.size = Pt(20)

        subtitle = document.add_paragraph("Generated by AgentFlow")
        subtitle.runs[0].font.color.rgb = RGBColor(71, 85, 105)
        subtitle.runs[0].font.size = Pt(11)

        metadata = document.add_paragraph()
        metadata.add_run(f"Task Type: {task.task_type}\n")
        metadata.add_run(f"Completed: {self._to_ist_string(task.completed_at)}\n")
        metadata.add_run(f"Task ID: {task.id}")
        metadata.runs[0].font.size = Pt(10)

        separator = document.add_paragraph("")
        self._add_docx_bottom_border(separator)

        content_text = output.content_text or ""
        for raw_line in self._parse_markdown_lines(content_text):
            line = raw_line.rstrip("\n")
            stripped = line.strip()

            if stripped.startswith("# "):
                heading = document.add_heading("", level=1)
                self._write_docx_markdown_runs(heading, stripped[2:].strip())
                continue
            if stripped.startswith("## "):
                heading = document.add_heading("", level=2)
                self._write_docx_markdown_runs(heading, stripped[3:].strip())
                continue
            if stripped.startswith("### "):
                heading = document.add_heading("", level=3)
                self._write_docx_markdown_runs(heading, stripped[4:].strip())
                continue
            if stripped.startswith("- ") or stripped.startswith("* "):
                paragraph = document.add_paragraph(style="List Bullet")
                self._write_docx_markdown_runs(paragraph, stripped[2:].strip())
                continue
            if self._is_numbered_list(stripped):
                numbered_text = re.sub(r"^\d+\.\s+", "", stripped)
                paragraph = document.add_paragraph(style="List Number")
                self._write_docx_markdown_runs(paragraph, numbered_text)
                continue
            if line.startswith("    ") or line.startswith("\t"):
                code_paragraph = document.add_paragraph(line, style="No Spacing")
                for run in code_paragraph.runs:
                    run.font.name = "Courier New"
                    run.font.size = Pt(10)
                continue
            if not stripped:
                document.add_paragraph("")
                continue

            paragraph = document.add_paragraph()
            self._write_docx_markdown_runs(paragraph, line)

        footer = document.sections[0].footer.paragraphs[0]
        footer.text = "AgentFlow Technologies Pvt. Ltd.    "
        footer.runs[0].font.size = Pt(9)
        footer.runs[0].font.color.rgb = RGBColor(100, 116, 139)
        self._add_docx_page_number(footer)

        buffer = BytesIO()
        document.save(buffer)
        buffer.seek(0)
        return buffer

    def export_to_pdf(self, task: AutomationTask, output: TaskOutput) -> BytesIO:
        """Export task output to branded PDF using reportlab."""

        buffer = BytesIO()
        doc = SimpleDocTemplate(
            buffer,
            pagesize=A4,
            rightMargin=0.75 * inch,
            leftMargin=0.75 * inch,
            topMargin=0.8 * inch,
            bottomMargin=0.8 * inch,
        )

        styles = getSampleStyleSheet()
        styles.add(
            ParagraphStyle(
                name="AFHeading1",
                parent=styles["Heading1"],
                fontName="Helvetica-Bold",
                fontSize=16,
                textColor=colors.HexColor("#0F172A"),
                spaceAfter=10,
            )
        )
        styles.add(
            ParagraphStyle(
                name="AFHeading2",
                parent=styles["Heading2"],
                fontName="Helvetica-Bold",
                fontSize=13,
                textColor=colors.HexColor("#1E40AF"),
                spaceAfter=8,
            )
        )
        styles.add(
            ParagraphStyle(
                name="AFHeading3",
                parent=styles["Heading3"],
                fontName="Helvetica-Bold",
                fontSize=11,
                textColor=colors.HexColor("#0369A1"),
                spaceAfter=6,
            )
        )
        styles.add(
            ParagraphStyle(
                name="AFCode",
                parent=styles["Normal"],
                fontName="Courier",
                fontSize=9,
                textColor=colors.HexColor("#111827"),
                backColor=colors.HexColor("#F3F4F6"),
                leftIndent=8,
                rightIndent=8,
                leading=12,
            )
        )

        story: list[Any] = []
        story.append(Paragraph(html.escape(task.task_name or "AgentFlow Output"), styles["Title"]))
        story.append(
            Paragraph(
                html.escape(
                    f"Generated by AgentFlow • {self._to_ist_string(task.completed_at)}"
                ),
                styles["Normal"],
            )
        )
        story.append(Spacer(1, 12))

        bullet_items: list[Any] = []

        for raw_line in self._parse_markdown_lines(output.content_text or ""):
            line = raw_line.rstrip("\n")
            stripped = line.strip()
            heading = self._parse_markdown_heading(stripped)

            if heading:
                if bullet_items:
                    story.append(ListFlowable(bullet_items, bulletType="bullet", start="circle"))
                    bullet_items = []
                level, heading_text = heading
                if level == 1:
                    style_name = "AFHeading1"
                elif level == 2:
                    style_name = "AFHeading2"
                else:
                    style_name = "AFHeading3"
                story.append(Paragraph(self._to_pdf_paragraph_markup(heading_text), styles[style_name]))
                continue

            if stripped.startswith("- ") or stripped.startswith("* "):
                bullet_items.append(
                    ListItem(Paragraph(self._to_pdf_paragraph_markup(stripped[2:].strip()), styles["Normal"]))
                )
                continue
            if line.startswith("    ") or line.startswith("\t"):
                if bullet_items:
                    story.append(ListFlowable(bullet_items, bulletType="bullet", start="circle"))
                    bullet_items = []
                story.append(Preformatted(line, styles["AFCode"]))
                continue
            if not stripped:
                if bullet_items:
                    story.append(ListFlowable(bullet_items, bulletType="bullet", start="circle"))
                    bullet_items = []
                story.append(Spacer(1, 8))
                continue

            if bullet_items:
                story.append(ListFlowable(bullet_items, bulletType="bullet", start="circle"))
                bullet_items = []
            story.append(Paragraph(self._to_pdf_paragraph_markup(line), styles["Normal"]))

        if bullet_items:
            story.append(ListFlowable(bullet_items, bulletType="bullet", start="circle"))

        story.append(Spacer(1, 18))
        story.append(
            Paragraph(
                "AgentFlow Technologies Pvt. Ltd.",
                ParagraphStyle(
                    "AFFooter",
                    parent=styles["Normal"],
                    fontSize=9,
                    textColor=colors.HexColor("#64748B"),
                    alignment=1,
                ),
            )
        )

        def _draw_footer(canvas_obj: Any, doc_obj: Any) -> None:
            canvas_obj.saveState()
            canvas_obj.setFont("Helvetica", 8)
            canvas_obj.setFillColor(colors.HexColor("#94A3B8"))
            canvas_obj.drawString(0.75 * inch, 0.5 * inch, "AgentFlow Technologies Pvt. Ltd.")
            canvas_obj.drawRightString(
                A4[0] - 0.75 * inch,
                0.5 * inch,
                f"Page {doc_obj.page}",
            )
            canvas_obj.restoreState()

        doc.build(story, onFirstPage=_draw_footer, onLaterPages=_draw_footer)
        buffer.seek(0)
        return buffer

    def export_to_csv(self, task: AutomationTask, output: TaskOutput) -> BytesIO:
        """Export task output as CSV with metadata fallback."""

        buffer = BytesIO()
        wrapper = io.TextIOWrapper(buffer, encoding="utf-8-sig", newline="")
        writer = csv.writer(wrapper)

        content_text = self._to_plain_markdown_text(output.content_text or "")
        lines = [line for line in content_text.splitlines() if line.strip()]

        is_table_like = False
        if output.output_type == "table" and lines:
            is_table_like = True
        elif lines:
            split_rows = [line.split(",") for line in lines]
            col_counts = {len(row) for row in split_rows}
            is_table_like = len(col_counts) == 1 and next(iter(col_counts), 0) > 1

        if is_table_like:
            for line in lines:
                writer.writerow([cell.strip() for cell in line.split(",")])
        else:
            writer.writerow(["Field", "Value"])
            writer.writerow(["Task Name", task.task_name or ""])
            writer.writerow(["Task Type", task.task_type])
            writer.writerow(["Completed At", self._to_ist_string(task.completed_at)])
            writer.writerow(["Task ID", str(task.id)])
            writer.writerow([])
            writer.writerow(["Output", content_text])

        wrapper.flush()
        wrapper.detach()
        buffer.seek(0)
        return buffer

    def export_to_json(self, task: AutomationTask, output: TaskOutput) -> BytesIO:
        """Export task output as structured JSON payload."""

        duration_seconds: int | None = None
        if task.started_at and task.completed_at:
            duration_seconds = int((task.completed_at - task.started_at).total_seconds())

        export_data = {
            "metadata": {
                "task_id": str(task.id),
                "task_name": task.task_name,
                "task_type": task.task_type,
                "status": task.status,
                "created_at": task.created_at.isoformat() if task.created_at else None,
                "completed_at": task.completed_at.isoformat() if task.completed_at else None,
                "duration_seconds": duration_seconds,
                "input": task.input_json,
                "output_type": output.output_type,
                "generated_by": "AgentFlow AI Platform",
                "platform": "agentflow.ai",
            },
            "output": {
                "content": output.content_text,
                "format": output.output_type,
                "word_count": len((output.content_text or "").split()),
                "character_count": len(output.content_text or ""),
            },
        }

        json_str = json.dumps(export_data, indent=2, ensure_ascii=False, default=str)
        buffer = BytesIO(json_str.encode("utf-8"))
        buffer.seek(0)
        return buffer

    def get_content_for_format(
        self,
        task: AutomationTask,
        output: TaskOutput,
        fmt: str,
    ) -> tuple[BytesIO, str, str]:
        """Dispatch export generation by format."""

        task_name_slug = self._task_name_slug(task.task_name or "agentflow_output")
        normalized = (fmt or "json").lower().strip()

        if normalized == "docx":
            return (
                self.export_to_docx(task, output),
                "application/vnd.openxmlformats-officedocument.wordprocessingml.document",
                f"{task_name_slug}.docx",
            )
        if normalized == "pdf":
            return (
                self.export_to_pdf(task, output),
                "application/pdf",
                f"{task_name_slug}.pdf",
            )
        if normalized == "csv":
            return (
                self.export_to_csv(task, output),
                "text/csv",
                f"{task_name_slug}.csv",
            )
        if normalized == "json":
            return (
                self.export_to_json(task, output),
                "application/json",
                f"{task_name_slug}.json",
            )

        raise ExportServiceError(f"Unsupported format: {fmt}")

    def tasks_to_csv(self, tasks: list[AutomationTask]) -> BytesIO:
        """Export task history rows to CSV buffer."""

        buffer = BytesIO()
        wrapper = io.TextIOWrapper(buffer, encoding="utf-8-sig", newline="")
        writer = csv.writer(wrapper)

        writer.writerow(
            [
                "Task Name",
                "Task Type",
                "Category",
                "Status",
                "Created At (IST)",
                "Completed At (IST)",
                "Duration (seconds)",
                "Error Message",
            ]
        )

        for task in tasks:
            category = ""
            try:
                from app.services.agent_runner import TASK_REGISTRY

                category = str(
                    TASK_REGISTRY.get(task.task_type, {}).get("category_display")
                    or TASK_REGISTRY.get(task.task_type, {}).get("category")
                    or task.task_type
                )
            except Exception:  # pylint: disable=broad-except
                category = task.task_type

            duration_seconds = ""
            if task.started_at and task.completed_at:
                duration_seconds = int((task.completed_at - task.started_at).total_seconds())

            writer.writerow(
                [
                    task.task_name or "",
                    task.task_type,
                    category,
                    task.status,
                    self._to_ist_string(task.created_at),
                    self._to_ist_string(task.completed_at),
                    duration_seconds,
                    task.error_message or "",
                ]
            )

        wrapper.flush()
        wrapper.detach()
        buffer.seek(0)
        return buffer


export_service = ExportService()
